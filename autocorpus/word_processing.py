"""This module provides functionality for processing supplementary files.

Extracts data from various file types such as PDFs, spreadsheets,
PowerPoint presentations, and archives. It also handles logging and error
management for unprocessed files.
"""

import datetime
import json
import os
import platform
import re
import subprocess
from pathlib import Path

import docx
import pandas as pd
from bioc import BioCCollection, BioCDocument, BioCPassage
from BioCTable import get_tables_bioc
from docx.document import Document

from autocorpus.utils import replace_unicode

from . import logger

WORD_EXTENSIONS = [".doc", ".docx"]


def extract_table_from_text(text: str) -> tuple[list[str], list[pd.DataFrame]]:
    """Extracts tables from a given text and returns the modified text and extracted tables.

    Args:
        text (str): The input text containing potential table data.

    Returns:
        tuple[str, list[pd.DataFrame]]: A tuple containing the modified text without table lines
        and a list of DataFrames representing the extracted tables.
    """
    # Split the text into lines
    lines = [x for x in text.splitlines() if x]
    text_output = lines

    # store extracted tables
    tables = []
    # Identify where the table starts and ends by looking for lines containing pipes
    table_lines = []
    # keep unmodified lines used in tables. These must be removed from the original text
    lines_to_remove = []
    inside_table = False
    for line in lines:
        if "|" in line:
            inside_table = True
            table_lines.append(line)
            lines_to_remove.append(line)
        elif (
            inside_table
        ):  # End of table if there's a blank line after lines with pipes
            inside_table = False
            tables.append(table_lines)
            table_lines = []
            continue

    for line in lines_to_remove:
        text_output.remove(line)

    tables_output = []
    # Remove lines that are just dashes (table separators)
    for table in tables:
        table = [line for line in table if not re.match(r"^\s*-+\s*$", line)]

        # Extract rows from the identified table lines
        rows = []
        for line in table:
            # Match only lines that look like table rows (contain pipes)
            if re.search(r"\|", line):
                # Split the line into cells using the pipe delimiter and strip whitespace
                cells = [
                    cell.strip()
                    for cell in line.split("|")
                    if not all(x in "|-" for x in cell)
                ]
                if cells:
                    # Remove empty cells that may result from leading/trailing pipes
                    # if cells[0] == '':
                    #     cells.pop(0)
                    # if cells[-1] == '':
                    #     cells.pop(-1)
                    rows.append(cells)

        # Determine the maximum number of columns in the table
        num_columns = max(len(row) for row in rows)

        # Pad rows with missing cells to ensure they all have the same length
        for row in rows:
            while len(row) < num_columns:
                row.append("")

        # Create a DataFrame from the rows
        df = pd.DataFrame(rows[1:], columns=rows[0])
        tables_output.append(df)
    return text_output, tables_output


def get_word_paragraphs(doc: Document) -> list[tuple[str, bool]]:
    """Extracts paragraphs from a Word document and identifies headers.

    Args:
        doc (Document): The Word document object.

    Returns:
        list[tuple[str, bool]]: A list of tuples where each tuple contains the paragraph text
        and a boolean indicating if it is a header.
    """
    text_sizes = set(
        [
            int(x.style.font.size)
            for x in doc.paragraphs
            if x.style and x.style.font and x.style.font.size
        ]
    )
    paragraphs = [
        (
            x.text,
            True
            if text_sizes
            and x.style
            and x.style.font
            and x.style.font.size
            and int(x.style.font.size) > min(text_sizes)
            else False,
        )
        for x in doc.paragraphs
    ]
    return paragraphs


def get_bioc_passages(text: list[str] | str) -> list[BioCPassage] | list[str]:
    """Identifies passages within the given text and creates passage objects.

    Args:
        text (list): The text to be processed, represented as a list of lines.

    Returns:
        list: A list of BioCPassage objects.
    """
    offset = 0
    passages: list[BioCPassage] = []
    if not text:
        return passages
    if isinstance(text, str):
        text = text.split("\n\n")
    text = [x for x in text if x]
    # Iterate through each line in the text
    for line in text:
        # Determine the type of the line and assign appropriate information
        iao_name = "supplementary material section"
        iao_id = "IAO:0000326"
        # Create a passage object and add it to the passages list
        passage = BioCPassage()
        passage.offset = offset
        passage.infons = {"iao_name_1": iao_name, "iao_id_1": iao_id}
        passage.text = line
        passages.append(passage)
        offset += len(line)
    return passages


def get_text_bioc(parsed_texts: list[str], filename: str):
    """Convert parsed texts into BioC format.

    Args:
        parsed_texts (list): A list of parsed text segments to be converted.
        filename (str): The name of the source file.
        textsource (str): The source of the text, default is "Auto-CORPus".

    Returns:
        BioCCollection: A BioCCollection object representing the converted text in BioC format.
    """
    passages = [
        p
        for sublist in [
            get_bioc_passages(replace_unicode(x)).__dict__["passages"]
            for x in parsed_texts
        ]
        for p in sublist
    ]
    offset = 0
    for p in passages:
        p["offset"] = offset
        offset += len(p["text"])
    # Create a BioC XML structure dictionary
    bioc = BioCCollection()
    bioc.source = "Auto-CORPus (supplementary)"
    bioc.date = datetime.date.today().strftime("%Y%m%d")
    bioc.key = "autocorpus_supplementary.key"
    bioc.documents = []
    new_doc = BioCDocument()
    new_doc.id = "1"
    new_doc.infons = {
        "inputfile": Path(filename).name,
        "textsource": "Auto-CORPus (supplementary)",
    }
    new_doc.passages = passages
    return bioc


class BioCText:
    def __init__(self, text):
        self.infons = {}
        self.passages = self.__identify_passages(text)
        self.annotations = []

    @staticmethod
    def __identify_passages(text):
        """Identifies passages within the given text and creates passage objects.

        Args:
            text (tuple): The text to be processed and a boolean which is True for header text.

        Returns:
            list: A list of passage objects. Each passage object is a dictionary containing the following keys:
                  - "offset": The offset of the passage in the original text.
                  - "infons": A dictionary of information associated with the passage, including:
                      - "iao_name_1": The name or type of the passage.
                      - "iao_id_1": The unique identifier associated with the passage.
                  - "text": The content of the passage.
                  - "sentences": An empty list of sentences (to be populated later if needed).
                  - "annotations": An empty list of annotations (to be populated later if needed).
                  - "relations": An empty list of relations (to be populated later if needed).

        Example:
            text = [
                "Introduction",
                "This is the first paragraph.",
                "Conclusion"
            ]
            passages = __identify_passages(text)
        """
        offset = 0
        passages = []
        # Iterate through each line in the text
        line, is_header = text
        line = line.replace("\n", "")
        iao_name = ""
        iao_id = ""

        # Determine the type of the line and assign appropriate information
        if line.isupper() or is_header:
            iao_name = "document title"
            iao_id = "IAO:0000305"
        else:
            iao_name = "supplementary material section"
            iao_id = "IAO:0000326"
        # Create a passage object and add it to the passages list
        passages.append(
            {
                "offset": offset,
                "infons": {"iao_name_1": iao_name, "iao_id_1": iao_id},
                "text": line,
                "sentences": [],
                "annotations": [],
                "relations": [],
            }
        )
        offset += len(line)
        return passages


class BioCTable:
    """Converts tables from nested lists into a BioC table object."""

    def __init__(self, input_file, table_id, table_data):
        self.inputfile = input_file
        self.id = str(table_id) + "_1"
        self.infons = {}
        self.passage = {}
        self.annotations = []
        self.__build_table(table_data)

    def __build_table(self, table_data):
        """Builds a table passage in a specific format and appends it to the list of passages.

        Args:
            table_data (list): The table data to be included in the passage. It should be a list
                               containing the table's column headings as the first row, followed by
                               the data rows.

        Returns:
            None

        Example:
            table_data = [
                ["Column 1", "Column 2", "Column 3"],
                [1, 2, 3],
                [4, 5, 6]
            ]
            self.__build_table(table_data)
        """
        passage = {
            "offset": 0,
            "infons": {
                "section_title_1": "table_content",
                "iao_name_1": "table",
                "iao_id_1": "IAO:0000306",
            },
            "column_headings": [],
            "data_section": [{"table_section_title_1": "", "data_rows": []}],
        }
        # Process the column headings of the table
        for i, col in enumerate(table_data[0]):
            passage["column_headings"].append(
                {"cell_id": self.id + f".1.{i + 1}", "cell_text": col}
            )
        # Process the data rows of the table
        for row_idx, row in enumerate(table_data[1:]):
            new_row = []
            for cell_idx, cell in enumerate(row):
                new_cell = {
                    "cell_id": f"{self.id}.{row_idx + 2}.{cell_idx + 1}",
                    "cell_text": f"{cell}",
                }
                new_row.append(new_cell)
            passage["data_section"][0]["data_rows"].append(new_row)
        self.passage = passage

    def get_table(self):
        return self.passage


def extract_tables(doc):
    """Extracts tables from a .docx document.

    Args:
        doc (docx.Document): The Document object representing the .docx document.

    Returns:
        list: A list of tables extracted from the document. Each table is represented as a nested list,
              where each inner list corresponds to a row, and each element in the inner list corresponds
              to the text content of a cell in the row.

    Example:
        from docx import Document

        doc = Document("document.docx")
        tables = extract_tables(doc)
    """
    # Open the .docx file
    tables = []
    # Iterate through the tables in the document
    for table in doc.tables:
        tables.append([])
        # Iterate through the rows in the table
        for row in table.rows:
            tables[-1].append([x.text for x in row.cells])
    return tables


def convert_older_doc_file(file: Path, output_dir: Path) -> Path | bool:
    """Converts an older .doc file to .docx format.

    Args:
        file (str): The path to the .doc file to be converted.
        output_dir (str): The directory where the converted .docx file will be saved.

    Returns:
        str: The path to the converted .docx file, or an empty string if the conversion fails.
    """
    operating_system = platform.system()
    docx_path = Path(str(file) + ".docx")
    if operating_system == "Windows":
        import win32com.client

        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            doc = word.Documents.Open(file)
            doc.SaveAs(docx_path, 16)
            doc.Close()
            word.Quit()
            return docx_path
        except Exception:
            return False
        finally:
            if word:
                word.Quit()
    elif operating_system == "linux":
        # Convert .doc to .docx using LibreOffice
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                output_dir,
                file,
            ],
            check=True,
            capture_output=True,
        )
        return docx_path
    elif operating_system == "Darwin":  # macOS
        try:
            # AppleScript to open the file in Word and save as .docx
            applescript = f"""
            tell application "Microsoft Word"
                open "{file}"
                save as active document file name "{docx_path}" file format format document
                close active document saving no
            end tell
            """
            subprocess.run(["osascript", "-e", applescript], check=True)
            return docx_path
        except Exception:
            return False
    else:
        return False


def extract_text_from_doc(
    file_path: Path,
) -> tuple[list[tuple[str, bool]], list[pd.DataFrame]]:
    """Extracts text and tables from a .doc file.

    Converts older .doc files to .docx format if necessary, then extracts text and tables.
    Deletes the temporary .docx file after processing.

    Args:
        file_path (str): The path to the .doc file.

    Returns:
        tuple: A tuple containing a list of paragraphs and a list of tables extracted from the document.

    Raises:
        ValueError: If the input file is not a .doc file.
        FileNotFoundError: If LibreOffice 'soffice' command is not found on Linux.
        Exception: For other errors during file processing.
    """
    if not file_path.suffix.lower() == ".doc":
        raise ValueError("Input file must be a .doc file.")
    try:
        output_dir = file_path.parent
        docx_path = convert_older_doc_file(file_path, output_dir)
        if isinstance(docx_path, Path):
            # Extract text from the resulting .docx file
            doc = docx.Document(str(docx_path.absolute()))
            tables = extract_tables(doc)
            paragraphs = get_word_paragraphs(doc)
            os.unlink(docx_path)
            return paragraphs, tables
        else:
            logger.info("Failed to convert .doc file to .docx.")
            return [], []
    except FileNotFoundError:
        print(
            "LibreOffice 'soffice' command not found. Ensure it is installed and in your PATH."
        )
        logger.warning(
            "LibreOffice 'soffice' command not found. Ensure it is installed and in your PATH."
        )
        return [], []
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {e}")
        return [], []


def process_word_document(file: Path, output_location: Path):
    """Processes a Word document to extract tables and paragraphs.

    Args:
        file (Path): The path to the Word document file to be processed.
        output_location (Path): The directory where the extracted data will be saved.

    Returns:
        bool: True if the document was processed successfully, False otherwise.
    """
    tables, paragraphs = [], []
    # Check if the file has a ".doc" or ".docx" extension
    if file.suffix.lower() in [".doc", ".docx"]:
        try:
            doc = docx.Document(str(file))
            tables = extract_tables(doc)
            paragraphs = get_word_paragraphs(doc)
        except ValueError:
            try:
                if not file.suffix.lower() == ".docx":
                    paragraphs, tables = extract_text_from_doc(file)
                    if paragraphs:
                        logger.info(
                            f"File {file} was converted to .docx as a copy within the same directory for processing."
                        )
                    else:
                        logger.info(
                            f"File {file} could not be processed correctly. It is likely a pre-2007 word document or problematic."
                        )
                        return False
                else:
                    logger.info(f"File {file} could not be processed correctly.")
                    return False
            except ValueError as ve:
                logger.info(f"File {file} raised the error:\n{ve}")
                return False
        except Exception as ex:
            logger.info(f"File {file} raised the error:\n{ex}")
            return False
    else:
        return False

    # Save tables as a JSON file
    if tables:
        if not output_location.exists():
            output_location.mkdir(parents=True, exist_ok=True)
        with open(f"{output_location}_tables.json", "w+", encoding="utf-8") as f_out:
            json.dump(get_tables_bioc(tables), f_out)

    # Save paragraphs as a JSON file
    if paragraphs:
        paragraphs = [x for x in paragraphs if x[0]]
        if not output_location.exists():
            output_location.mkdir(parents=True, exist_ok=True)
        with open(f"{output_location}_bioc.json", "w+", encoding="utf-8") as f_out:
            # TODO: Test if datatype causes a problem
            text = get_text_bioc(paragraphs, Path(file).name)
            json.dump(text, f_out, indent=4)

    if not paragraphs and not tables:
        return False
    else:
        return True
