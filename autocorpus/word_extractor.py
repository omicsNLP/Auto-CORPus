"""This module provides functionality to extract text and tables from Word documents (.doc and .docx).

It includes methods to handle older .doc files by converting them to .docx format and processing them.
"""

import os
import platform
import subprocess
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from pandas import DataFrame

from . import logger
from .ac_bioc.bioctable.collection import BioCTableCollection
from .ac_bioc.bioctable.json import BioCTableJSON
from .ac_bioc.collection import BioCCollection
from .ac_bioc.json import BioCJSON
from .bioc_supplementary import (
    BioCTableConverter,
    BioCTextConverter,
    WordText,
)


def __extract_tables(doc: DocumentObject) -> list[DataFrame]:
    """Extracts tables from a .docx document as a list of DataFrames.

    Args:
        doc: The Document object representing the .docx document.

    Returns:
        List[pd.DataFrame]: A list of pandas DataFrames, each representing a table in the document.

    Example:
        from docx import Document

        doc = Document("document.docx")
        tables = __extract_tables(doc)
    """
    dataframes: list[DataFrame] = []

    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        df = DataFrame(data)
        dataframes.append(df)

    return dataframes


def __windows_convert_doc_to_docx(docx_path: Path, file: Path) -> Path | None:
    """Converts a .doc file to .docx format using Microsoft Word on Windows."""
    try:
        import win32com.client
    except ImportError:
        logger.error(
            "pywin32 is required to convert Word documents on Windows. Please install it via 'pip install pywin32'."
        )
        return None

    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        doc = word.Documents.Open(str(file))
        doc.SaveAs(str(docx_path), 16)  # 16 = wdFormatDocumentDefault (.docx)
        doc.Close()
        logger.info(
            f"Successfully converted '{file}' to '{docx_path}' using Word on Windows."
        )
        return docx_path
    except Exception as e:
        logger.exception(f"Failed to convert '{file}' on Windows: {e}")
        return None
    finally:
        if word:
            try:
                word.Quit()
            except Exception as quit_err:
                logger.warning(f"Could not quit Word application cleanly: {quit_err}")


def __linux_convert_doc_to_docx(docx_path: Path, file: Path) -> Path | None:
    """Converts a .doc file to .docx format using LibreOffice on Linux."""
    try:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(docx_path.parent),
                str(file),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        logger.info(f"LibreOffice output: {result.stdout}")
        return docx_path
    except FileNotFoundError:
        logger.error(
            "LibreOffice ('soffice') not found. Please install it to enable DOC to DOCX conversion."
        )
        return None
    except subprocess.CalledProcessError as e:
        logger.exception(f"LibreOffice failed to convert '{file}': {e.stderr}")
        return None


def __escape_applescript_path(path: Path) -> str:
    # Convert to absolute path just in case
    path = path.absolute()
    # Escape backslashes and double quotes for AppleScript
    return str(path).replace("\\", "\\\\").replace('"', '\\"')


def __macos_convert_doc_to_docx(docx_path: Path, file: Path) -> Path | None:
    """Converts a .doc file to .docx format using AppleScript on macOS."""
    try:
        applescript = f'''
        tell application "Microsoft Word"
            open "{__escape_applescript_path(file)}"
            save as active document file name "{__escape_applescript_path(docx_path)}" file format format document
            close active document saving no
        end tell
        '''
        subprocess.run(["osascript", "-e", applescript], check=True)
        logger.info(
            f"Successfully converted '{file}' to '{docx_path}' using Word on macOS."
        )
        return docx_path
    except FileNotFoundError:
        logger.error(
            "osascript not found. Ensure you have AppleScript and Microsoft Word installed on macOS."
        )
        return None
    except subprocess.CalledProcessError as e:
        logger.exception(f"AppleScript failed to convert '{file}': {e}")
        return None


def __convert_older_doc_file(file: Path, output_dir: Path) -> Path | None:
    """Converts an older .doc file to .docx format using platform-specific methods."""
    operating_system = platform.system()
    docx_path = output_dir / file.with_suffix(".docx").name

    if operating_system == "Windows":
        return __windows_convert_doc_to_docx(docx_path, file)
    elif operating_system == "Darwin":  # macOS
        return __macos_convert_doc_to_docx(docx_path, file)
    else:
        return __linux_convert_doc_to_docx(docx_path, file)  # Fallback to Linux method


def extract_word_content(file_path: Path):
    """Extracts text from a .doc file by converting it to .docx and processing with python-docx."""
    if file_path.suffix.lower() not in [".doc", ".docx"]:
        raise ValueError("Input file must be a .doc file.")
    try:
        output_dir = Path(file_path).parent.absolute()
        # Check if the file is a .doc file
        if file_path.suffix.lower() == ".doc":
            docx_path = __convert_older_doc_file(file_path, output_dir)

        # Extract text from the resulting .docx file
        doc = Document(str(docx_path))
        tables = __extract_tables(doc)
        text_sizes = set(
            [
                int(x.style.font.size)
                for x in doc.paragraphs
                if x.style and x.style.font.size
            ]
        )
        paragraphs = [
            WordText(
                x.text,
                True
                if text_sizes
                and x.style
                and x.style.font.size
                and int(x.style.font.size) > min(text_sizes)
                else False,
            )
            for x in doc.paragraphs
        ]
        bioc_text: BioCCollection | None = None
        bioc_tables: BioCTableCollection | None = None

        if paragraphs:
            bioc_text = BioCTextConverter.build_bioc(paragraphs, str(file_path), "word")

        if tables:
            bioc_tables = BioCTableConverter.build_bioc(tables, str(file_path))

        if bioc_text:
            out_filename = str(file_path).replace(
                file_path.suffix, f"{file_path.suffix}_bioc.json"
            )
            with open(out_filename, "w", encoding="utf-8") as f:
                BioCJSON.dump(bioc_text, f, indent=4)

        if bioc_tables:
            out_table_filename = str(file_path).replace(
                file_path.suffix, f"{file_path.suffix}_tables.json"
            )
            with open(out_table_filename, "w", encoding="utf-8") as f:
                BioCTableJSON.dump(bioc_tables, f, indent=4)

        os.unlink(str(docx_path))
    except FileNotFoundError:
        logger.error(
            "LibreOffice 'soffice' command not found. Ensure it is installed and in your PATH."
        )
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {e}")
