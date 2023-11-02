import datetime
import json
import os
import platform
import subprocess
from os.path import join
import logging
from docx import Document

logging.basicConfig(filename="WordExtractor.log", level=logging.ERROR, format="%(asctime)s - %(levelname)s - %("
                                                                              "message)s")

filename = ""


class BioCText:
    def __init__(self, input_file, text):
        self.inputfile = input_file
        self.infons = {}
        self.passages = self.__identify_passages(text)
        self.annotations = []

    @staticmethod
    def __identify_passages(text):
        """
        Identifies passages within the given text and creates passage objects.

        Args:
            text (list): The text to be processed, represented as a list of lines.

        Returns:
            list: A list of passage objects. Each passage object is a dictionary containing the following keys:
                  - "offset": The offset of the passage in the original text.
                  - "infons": A dictionary of information associated with the passage, including:
                      - "iao_name_1": The name or type of the passage.
                      - "iao_id_1": The unique identifier associated with the passage.
                  - "text": The content of the passage.
                  - "sentences": An empty list of sentences (to be populated later if needed).
                  - "annotations": An empty list of annotations (to be populated later if needed).
                  - "relations": An empty list of relations (to be populated later if needed).

        Example:
            text = [
                "Introduction",
                "This is the first paragraph.",
                "Conclusion"
            ]
            passages = __identify_passages(text)
        """
        offset = 0
        passages = []
        # Iterate through each line in the text
        for line, is_header in text:
            line = line.replace("\n", "")
            iao_name = ""
            iao_id = ""

            # Determine the type of the line and assign appropriate information
            if line.isupper() or is_header:
                iao_name = "document title"
                iao_id = "IAO:0000305"
            else:
                iao_name = "supplementary material section"
                iao_id = "IAO:0000326"
            # Create a passage object and add it to the passages list
            passages.append({
                "offset": offset,
                "infons": {
                    "iao_name_1": iao_name,
                    "iao_id_1": iao_id
                },
                "text": line,
                "sentences": [],
                "annotations": [],
                "relations": []
            })
            offset += len(line)
        return passages


class BioCTable:
    """
    Converts tables from nested lists into a BioC table object.
    """

    def __init__(self, input_file, table_id, table_data):
        self.inputfile = input_file
        self.id = str(table_id) + "_1"
        self.infons = {}
        self.passages = []
        self.annotations = []
        self.__build_table(table_data)

    def __build_table(self, table_data):
        """
        Builds a table passage in a specific format and appends it to the list of passages.

        Args:
            table_data (list): The table data to be included in the passage. It should be a list
                               containing the table's column headings as the first row, followed by
                               the data rows.

        Returns:
            None

        Example:
            table_data = [
                ["Column 1", "Column 2", "Column 3"],
                [1, 2, 3],
                [4, 5, 6]
            ]
            self.__build_table(table_data)
        """
        passage = {
            "offset": 0,
            "infons": {
                "section_title_1": "table_content",
                "iao_name_1": "table",
                "iao_id_1": "IAO:0000306"
            },
            "column_headings": [],
            "data_section": [
                {
                    "table_section_title_1": "",
                    "data_rows": [

                    ]
                }
            ]
        }
        # Process the column headings of the table
        for i, col in enumerate(table_data[0]):
            passage["column_headings"].append(
                {
                    "cell_id": self.id + F".1.{i + 1}",
                    "cell_text": col
                }
            )
        # Process the data rows of the table
        for row_idx, row in enumerate(table_data[1:]):
            new_row = []
            for cell_idx, cell in enumerate(row):
                new_cell = {
                    "cell_id": F"{self.id}.{row_idx + 2}.{cell_idx + 1}",
                    "cell_text": F"{cell}"
                }
                new_row.append(new_cell)
            passage["data_section"][0]["data_rows"].append(new_row)
        self.passages.append(passage)


def get_tables_bioc(tables):
    """
    Generates a BioC XML structure containing tables.

    Args:
        tables (list): A list of tables to be included in the BioC structure.
                       Each table should be represented as a nested list, where each inner list
                       corresponds to a row, and each element in the inner list corresponds to the
                       text content of a cell in the row.

    Returns:
        dict: A dictionary representing the generated BioC XML structure.

    Example:
        tables = [[["A", "B"], ["1", "2"]], [["X", "Y"], ["3", "4"]]]
        bioc_xml = get_tables_bioc(tables)
    """
    global filename
    # Create a BioC XML structure dictionary
    bioc = {
        "source": "Auto-CORPus (supplementary)",
        "date": str(datetime.date.today().strftime("%Y%m%d")),
        "key": "autocorpus_supplementary.key",
        "infons": {},
        "documents": [
            {
                "id": 1,
                "inputfile": filename,
                "infons": {},
                "passages": [BioCTable(filename, i + 1, x).__dict__ for i, x in enumerate(tables)],
                "annotations": [],
                "relations": []
            }
        ]
    }
    return bioc


def get_text_bioc(paragraphs):
    """
    Generates a BioC XML structure containing text paragraphs.

    Args:
        paragraphs (list): A list of paragraphs to be included in the BioC structure.

    Returns:
        dict: A dictionary representing the generated BioC XML structure.

    Example:
        paragraphs = ["This is the first paragraph.", "This is the second paragraph."]
        bioc_xml = get_text_bioc(paragraphs)
    """
    global filename
    # Create a BioC XML structure dictionary
    bioc = {
        "source": "Auto-CORPus (supplementary)",
        "date": str(datetime.date.today().strftime("%Y%m%d")),
        "key": "autocorpus_supplementary.key",
        "infons": {},
        "documents": [
            {
                "id": 1,
                "inputfile": filename,
                "infons": {},
                "passages": [BioCText(filename, paragraphs).__dict__],
                "annotations": [],
                "relations": []
            }
        ]
    }
    return bioc


def extract_tables(doc):
    """
    Extracts tables from a .docx document.

    Args:
        doc (docx.Document): The Document object representing the .docx document.

    Returns:
        list: A list of tables extracted from the document. Each table is represented as a nested list,
              where each inner list corresponds to a row, and each element in the inner list corresponds
              to the text content of a cell in the row.

    Example:
        from docx import Document

        doc = Document("document.docx")
        tables = extract_tables(doc)
    """
    # Open the .docx file
    tables = []
    # Iterate through the tables in the document
    for table in doc.tables:
        tables.append([])
        # Iterate through the rows in the table
        for row in table.rows:
            tables[-1].append([x.text for x in row.cells])
    return tables


def convert_older_doc_file(file):
    operating_system = platform.system()
    if operating_system == "Windows":
        import win32com.client
        try:
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(file)
            doc.SaveAs(file + ".docx", 16)
            return True
        except Exception as e:
            return False
    elif operating_system == "linux":
        subprocess.call(['unoconv', '-d', 'document', '--format=docx', file])
        return True
    elif operating_system == "mac":
        return False


def process_word_document(file):
    """
    Processes a Word document file, extracting tables and paragraphs, and saving them as JSON files.

    Args:
        file (str): The path to the Word document file.

    Returns:
        None

    Example:
        file_path = "/path/to/document.docx"
        process_word_document(file_path)
    """
    tables, paragraphs = [], []
    # Check if the file has a ".doc" or ".docx" extension
    if file.endswith(".doc") or file.endswith(".docx"):
        try:
            doc = Document(file)
            tables = extract_tables(doc)
            text_sizes = set([int(x.style.font.size) for x in doc.paragraphs if x.style.font.size])
            paragraphs = [(x.text, True if text_sizes and x.style.font.size and int(x.style.font.size) > min(
                text_sizes) else False) for x in doc.paragraphs]
        except ValueError:
            if not file.endswith(".docx"):
                conversion_check = convert_older_doc_file(file)
                if conversion_check:
                    logging.info(
                        F"File {file} was converted to .docx as a copy within the same directory for processing.")
                    process_word_document(file + ".docx")
                else:
                    logging.info(
                        F"File {file} could not be processed correctly. It is likely a pre-2007 word document or problematic.")
            else:
                logging.info(F"File {file} could not be processed correctly.")
        except Exception as ex:
            logging.info(F"File {file} raised the error:\n{ex}")
    else:
        return

    # Save tables as a JSON file
    with open(F"{file}_tables_bioc.json", "w+", encoding="utf-8") as f_out:
        json.dump(get_tables_bioc(tables), f_out)

    # Save paragraphs as a JSON file
    with open(F"{file}_bioc.json", "w+", encoding="utf-8") as f_out:
        json.dump(get_text_bioc(paragraphs), f_out)


def process_directories(input_directory):
    """
        Processes Word documents in the specified input directory.

    Args:
        input_directory (str): The path to the input directory containing the Word documents.

    Returns:
        None

    Example:
        input_directory = "/path/to/documents"
        process_directories(input_directory)
    """
    global filename
    # Iterate over the files in the input directory and its subdirectories
    for parent, folders_in_parent, files_in_parent in os.walk(input_directory):
        for file in files_in_parent:
            # Check if the file has a ".doc" or ".docx" extension
            if file.endswith(".doc") or file.endswith(".docx"):
                filename = file
                # Process the Word document
                process_word_document(join(parent, file))
